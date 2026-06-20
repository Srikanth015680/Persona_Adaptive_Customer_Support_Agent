from dataclasses import dataclass
from pathlib import Path
from typing import List

from pypdf import PdfReader

from . import config


@dataclass
class RawDocument:
    source: str
    text: str


@dataclass
class Chunk:
    text: str
    source: str
    chunk_index: int
    section: str


def load_documents(data_dir: Path = None) -> List[RawDocument]:
    data_dir = data_dir or config.DATA_DIR
    docs = []

    for path in sorted(Path(data_dir).iterdir()):
        if path.name.startswith("_") or path.is_dir():
            continue

        suffix = path.suffix.lower()

        if suffix in (".md", ".txt"):
            text = path.read_text(encoding="utf-8")
        elif suffix == ".pdf":
            text = _extract_pdf_text(path)
        elif suffix == ".docx":
            text = _extract_docx_text(path)
        else:
            continue

        if text.strip():
            docs.append(RawDocument(source=path.name, text=text))

    return docs


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        parts.append(f"[Page {i + 1}]\n{page_text}")

    return "\n\n".join(parts)


def _extract_docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)


_SEPARATORS = ["\n\n", "\n", ". ", " ", ""]


def _split_text(text: str, chunk_size: int, chunk_overlap: int, separators=None) -> List[str]:
    separators = separators or _SEPARATORS

    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    sep = separators[0]
    remaining_separators = separators[1:]

    if sep == "":
        pieces = [
            text[i:i + chunk_size]
            for i in range(0, len(text), chunk_size - chunk_overlap)
        ]
        return [p for p in pieces if p.strip()]

    parts = text.split(sep)
    chunks = []
    current = ""

    for part in parts:
        candidate = (current + sep + part) if current else part

        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)

            if len(part) > chunk_size:
                chunks.extend(
                    _split_text(
                        part,
                        chunk_size,
                        chunk_overlap,
                        remaining_separators,
                    )
                )
                current = ""
            else:
                current = part

    if current:
        chunks.append(current)

    if chunk_overlap > 0 and len(chunks) > 1:
        overlapped = [chunks[0]]

        for i in range(1, len(chunks)):
            tail = overlapped[-1][-chunk_overlap:]
            overlapped.append((tail + " " + chunks[i]).strip())

        return overlapped

    return chunks


def _guess_section(chunk_text: str, doc_text: str) -> str:
    idx = doc_text.find(chunk_text[:40])

    if idx == -1:
        return "General"

    preceding = doc_text[:idx]

    headings = [
        line.lstrip("#").strip()
        for line in preceding.splitlines()
        if line.strip().startswith("#")
    ]

    return headings[-1] if headings else "General"


def chunk_documents(
    docs: List[RawDocument],
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> List[Chunk]:

    chunk_size = chunk_size or config.CHUNK_SIZE
    chunk_overlap = chunk_overlap or config.CHUNK_OVERLAP

    all_chunks = []

    for doc in docs:
        pieces = _split_text(doc.text, chunk_size, chunk_overlap)

        for idx, piece in enumerate(pieces):
            section = _guess_section(piece, doc.text)

            all_chunks.append(
                Chunk(
                    text=piece.strip(),
                    source=doc.source,
                    chunk_index=idx,
                    section=section,
                )
            )

    return all_chunks